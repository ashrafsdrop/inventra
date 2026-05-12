import sys
import subprocess
import os

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Add Title
title = doc.add_heading('Web Engineering: Enterprise Resource Planning (ERP) System', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('1. Introduction', level=1)
doc.add_paragraph("This part introduces the topic and sets the foundation for the report. It explains what web engineering is and why it's important, especially in the world of enterprise resource planning and business management.")

doc.add_heading('1.1 Overview of Web Engineering', level=2)
doc.add_paragraph("Web Engineering is the process of building and maintaining web-based applications in a proper, systematic way. It uses ideas from software engineering, information systems, and how people interact with computers. Web engineering ensures that applications are methodical, well-structured, and satisfy the objectives of both businesses and users by applying engineering concepts. The goal is to make web systems that are reliable, efficient, scalable, and easy to use.")

doc.add_heading('1.2 Role of Web Engineering in Modern Applications', level=2)
doc.add_paragraph("Web Engineering is crucial to modern technology as it drives digital transformation in almost every sector. Web engineering helps make sure that websites are high-quality, secure, scalable, and easy to maintain. In business, it powers ERP systems, customer service systems, and stock management. In education, it facilitates online classrooms and digital learning platforms. In healthcare, it enables virtual consultations, electronic records, and appointment scheduling systems. These examples demonstrate how Web Engineering improves daily business operations by enabling essential services.")

doc.add_heading('1.3 Importance of ERP Systems in Web Environment', level=2)
doc.add_paragraph("Global business operations have changed as a result of cloud computing and web technologies. It makes it possible for business owners, managers, and employees to access vital company data and manage operations from any place, at any time. This gives companies the ability to streamline their workflows, connect decentralized branches, and make data-driven decisions instantly. Web-based ERP is rapidly growing and now accounts for a sizable portion of global business management tools, replacing outdated on-premise solutions.")

doc.add_heading('1.4 Objectives of the Report', level=2)
doc.add_paragraph("This report's main goal is to analyze the foundational ideas of web engineering and show how to use them to create a safe, scalable, and reliable ERP platform named \"Inventra\". The paper will outline how to put these ideas into practice in order to create a system that meets user needs, handles complex data relationships, and operates efficiently.")

doc.add_heading('1.5 Scope of the ERP System', level=2)
doc.add_paragraph("This study will focus on the development of an ERP system with specific features. This includes user authentication (signing in), dashboard analytics, managing inventory and tracking stock levels, creating sales and purchase invoices, managing contacts (customers and suppliers), and recording financial transactions in a general ledger. The system will facilitate these main functions to automate day-to-day business operations.")

doc.add_heading('1.6 Methodology', level=2)
doc.add_paragraph("To prepare this report, a literature review was done. To understand the topic, this entails reading academic papers, industrial research, and technical literature. Mention the sources you used, including technical documentation (Next.js and Django REST Framework), industry reports, and scholarly articles. Additionally, the report assesses existing ERP systems to determine their benefits and drawbacks.")

doc.add_heading('2. Problem Statement', level=1)

doc.add_heading('2.1 Limitations of Traditional Retail & Business Models', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run("Restricted geographic reach and operational silos: ").bold = True
p.add_run("Operations are constrained to physical office locations or local servers.")

p = doc.add_paragraph(style='List Bullet')
p.add_run("High operating cost: ").bold = True
p.add_run("High operating expenses include manual inventory control, redundant data entry, physical paper trails, and excessive employee compensation for administrative tasks.")

p = doc.add_paragraph(style='List Bullet')
p.add_run("Limited working hours: ").bold = True
p.add_run("Unable to access business data outside of standard office hours. Lack of centralization makes it challenging to monitor company performance in real-time. Manual procedures like stock management, billing, and record-keeping take a lot of time and are prone to mistakes.")

doc.add_heading('2.2 Challenges in Existing Web-based Platforms', level=2)
p = doc.add_paragraph(style='List Bullet')
p.add_run("Performance issues: ").bold = True
p.add_run("Slow website speed, crashes during high data queries, poor user experience.")

p = doc.add_paragraph(style='List Bullet')
p.add_run("Scalability: ").bold = True
p.add_run("Difficulty in handling large numbers of users, products, and financial transactions as the business grows.")

p = doc.add_paragraph(style='List Bullet')
p.add_run("Security and privacy concerns: ").bold = True
p.add_run("Risk of hacking, data breaches, identity theft, and misuse of sensitive corporate financial data.")

p = doc.add_paragraph(style='List Bullet')
p.add_run("Data reliability: ").bold = True
p.add_run("Failing to keep inventory and accounting ledgers synchronized in real-time, leading to disastrous miscalculations.")

doc.add_heading('2.3 Objectives of the Proposed Web-based ERP Solution', level=2)
doc.add_paragraph("Create a dependable, quick, and easy-to-use internet-based enterprise management platform (Inventra).", style='List Number')
doc.add_paragraph("Make sure it is scalable to accommodate future increases in transactions, inventory size, and traffic.", style='List Number')
doc.add_paragraph("Put robust security mechanisms in place (HTTPS, SSL, encryption, JWT authentication).", style='List Number')
doc.add_paragraph("Include a number of trustworthy modules (Inventory, Sales, Purchases, Accounts) that seamlessly communicate with each other.", style='List Number')

output_path = os.path.join(os.getcwd(), 'Inventra_Project_Report.docx')
doc.save(output_path)
print(f"Successfully generated Word document at: {output_path}")
